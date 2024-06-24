# views.py
from django import views 
from django.shortcuts import render
from django.http import HttpResponse
from .forms import TextInputForm, TupleForm, TupleFormSet
import unicodedata
import time 
from django.views.generic import (
    CreateView,

)
from django.contrib.auth.decorators import user_passes_test
from django.shortcuts import render, redirect
import math 
from killgpt.users.models import User

    # Input all your text into this. The input() function is not used because it can only take a certain number of words. Also you can put it into hte input string 
# The function cannot take in quotation marks as inputs. So we parse them out with the remove_quotation_marks 

from django.shortcuts import render
from .forms import TextInputForm
import PyPDF2
import docx
import magic
import unicodedata
import time

import fitz

def process_file(file):
    file_type = magic.from_buffer(file.read(), mime=True)
    file.seek(0)
    if file_type == 'application/pdf':
        doc = fitz.open(stream=file.read(), filetype='pdf')
        text = ''
        for page in doc:
            text += page.get_text()
        return text
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = docx.Document(file)
        text = ''
        for para in doc.paragraphs:
            text += para.text
        return text
    elif file_type == 'text/plain':
        # Read and return the contents of a .txt file
        text = file.read().decode('utf-8')
        return text
    else:
        raise ValueError('Unsupported file type')



from .davebot import DaveBot
from django.http import FileResponse, HttpResponseRedirect
import os 
from io import StringIO

# Only signed in users can go onto the detector 
# @user_passes_test(lambda u: u.is_authenticated, login_url='/accounts/login/')
def TextInputView(request):
     # Retrieve the user object from the request or context
    
    if request.method == 'POST':
        if request.method == 'POST' and 'form-TOTAL_FORMS' in request.POST:
            formset_data = {}
            total_forms = int(request.POST.get('form-TOTAL_FORMS', 0))
            print(total_forms)
            for i in range(total_forms):
                formset_data[f'element_1_{i}'] = request.POST.get(f'form-{i}-element_1', '')
                formset_data[f'element_2_{i}'] = request.POST.get(f'form-{i}-element_2', '')

            formset_values = []
            for i in range(total_forms):
                formset_values.append({
                    'element_1': formset_data[f'element_1_{i}'],
                    'element_2': formset_data[f'element_2_{i}']
                })
            
            # Get the index of the last form with the overall feedback 
            last_form_index = total_forms - 1
            overall_feedback = request.POST.get(f'form-{last_form_index}-overall_feedback', '')
            print("Element 3:", overall_feedback)
            # Store formset_values and overall_feedback in session
            request.session['formset_values'] = formset_values
            request.session['overall_feedback'] = overall_feedback

            # Redirect to the download view
            html_content = loader.render_to_string('download_template.html', {'formset_values': formset_values, 'overall_feedback': overall_feedback})
            response = HttpResponse(html_content)
            return response
        
        else:
            form = TextInputForm(request.POST, request.FILES)
            if form.is_valid():
                initial = time.time()
                text_input = form.cleaned_data['text_input']
                rubric_criteria = form.cleaned_data['rubric_input']
                subject_input = form.cleaned_data['subject_input']
                file_input = form.cleaned_data['file_input']
               
                # validate the file's content
                if file_input:
                    try:
                        input_text = process_file(file_input)
                    except ValueError as e:
                        context = {'form': form, 'error': str(e)} 
                        return render(request, 'pages/home.html', context)
                else:
                    input_text = text_input
            
                
                student_writing = input_text
                

                dave = DaveBot()
                feedback_dict_sentences = dave.get_feedback(student_writing, rubric_criteria, subject_input)
                # Extract data from feedback_dict_sentences to pass as initial data to the formset
                initial_data = []
                for key, value in feedback_dict_sentences.items():
                    initial_data.append({'element_1': value[0], 'element_2': value[1]})  # Assuming each value is a tuple with two elements
                initial_data.append({'overall_feedback': "Please provide your overall feedback here.",})
                formset = TupleFormSet(initial=initial_data)
                context = {'form': form, 'input_text': input_text, 'formset': formset, 'feedback_dict_sentences': feedback_dict_sentences,}
                
            else:
                context = {'form': form}
    else:
        form = TextInputForm()
        context = {'form': form}
    return render(request, 'pages/home.html', context)

from django.shortcuts import redirect
import os
from django.template import loader

from docx import Document

def download_file(request):
    # Retrieve formset_values from session
    formset_values = request.session.get('formset_values', [])
    overall_f = request.session.get('overall_feedback')

    # Create a new Word document
    document = docx.Document()

    # Add sentences and feedback as paragraphs with comments
    for item in formset_values:
        paragraph = document.add_paragraph(item['element_1'])
        paragraph.add_comment(item['element_2'], author='Instructor')

    # Add overall feedback as a separate paragraph
    if overall_f:
        document.add_paragraph('Overall Feedback:')
        document.add_paragraph(overall_f)

    # Create an in-memory stream to store the Word document
    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)

    # Set response to download the Word file
    response = HttpResponse(stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="commented_file.docx"'

    return response
    # # Alternatively, you can render an HTML response before download
    # html_content = loader.render_to_string('download_template.html', {'formset_values': formset_values})
    # response = HttpResponse(html_content)
    

    
